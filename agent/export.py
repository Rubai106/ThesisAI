"""Export utilities for ThesisAI — PDF reading, Word/LaTeX export, file saving."""

import os
import re
from datetime import datetime
from typing import Optional

try:
    import fitz  # PyMuPDF
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


DRAFTS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "drafts")


def ensure_drafts_dir() -> str:
    """Create drafts directory if it doesn't exist."""
    os.makedirs(DRAFTS_DIR, exist_ok=True)
    return DRAFTS_DIR


# ══════════════════════════════════════════════════════════════════
#  PDF READING
# ══════════════════════════════════════════════════════════════════

def extract_pdf_text(filepath: str, max_pages: int = 50) -> str:
    """Extract text from a PDF file using PyMuPDF.

    Returns the full text content, truncated if very long.
    """
    if not HAS_FITZ:
        return "[Error: PyMuPDF is not available in this environment. PDF reading is disabled.]"
    doc = fitz.open(filepath)
    pages = []
    for i, page in enumerate(doc):
        if i >= max_pages:
            pages.append(f"\n\n[...truncated after {max_pages} pages]")
            break
        pages.append(page.get_text())
    doc.close()

    text = "\n\n".join(pages)
    # Truncate to ~15000 chars to stay within context limits
    if len(text) > 15000:
        text = text[:15000] + "\n\n[...truncated for length]"
    return text


# ══════════════════════════════════════════════════════════════════
#  SAVE DRAFT TO FILE
# ══════════════════════════════════════════════════════════════════

def save_draft_markdown(content: str, filename: str = "") -> str:
    """Save a draft as a Markdown file in the drafts folder.

    Returns the full path of the saved file.
    """
    ensure_drafts_dir()
    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"draft_{timestamp}.md"
    if not filename.endswith(".md"):
        filename += ".md"

    filepath = os.path.join(DRAFTS_DIR, _sanitize_filename(filename))
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)
    return filepath


# ══════════════════════════════════════════════════════════════════
#  EXPORT TO WORD (.docx)
# ══════════════════════════════════════════════════════════════════

def export_to_word(content: str, filename: str = "", title: str = "") -> str:
    """Convert markdown content to a Word document.

    Returns the full path of the saved .docx file.
    """
    if not HAS_DOCX:
        return "[Error: python-docx is not available in this environment. Word export is disabled.]"
    ensure_drafts_dir()
    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"thesis_{timestamp}.docx"
    if not filename.endswith(".docx"):
        filename += ".docx"

    filepath = os.path.join(DRAFTS_DIR, _sanitize_filename(filename))

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title
    if title:
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse markdown-ish content into Word
    _markdown_to_docx(doc, content)

    doc.save(filepath)
    return filepath


def _markdown_to_docx(doc: Document, text: str) -> None:
    """Simple markdown-to-docx converter for thesis content."""
    lines = text.split("\n")
    in_code_block = False
    code_buffer = []

    for line in lines:
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            if in_code_block:
                # End code block
                code_text = "\n".join(code_buffer)
                para = doc.add_paragraph()
                run = para.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        # Empty lines
        if not stripped:
            continue

        # Headings
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text_content = stripped[2:]
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(para, text_content)
        # Numbered lists
        elif re.match(r"^\d+\.\s", stripped):
            text_content = re.sub(r"^\d+\.\s", "", stripped)
            para = doc.add_paragraph(style="List Number")
            _add_formatted_text(para, text_content)
        # Regular paragraph
        else:
            para = doc.add_paragraph()
            _add_formatted_text(para, stripped)


def _add_formatted_text(para, text: str) -> None:
    """Add text with basic bold/italic formatting to a paragraph."""
    # Split on bold (**text**) and italic (*text*) markers
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            para.add_run(part)


# ══════════════════════════════════════════════════════════════════
#  EXPORT TO LaTeX
# ══════════════════════════════════════════════════════════════════

def export_to_latex(content: str, filename: str = "", title: str = "",
                    author: str = "Student Name") -> str:
    """Convert markdown content to a LaTeX document.

    Returns the full path of the saved .tex file.
    """
    ensure_drafts_dir()
    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"thesis_{timestamp}.tex"
    if not filename.endswith(".tex"):
        filename += ".tex"

    filepath = os.path.join(DRAFTS_DIR, _sanitize_filename(filename))

    latex = _markdown_to_latex(content, title=title, author=author)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(latex)
    return filepath


def _markdown_to_latex(text: str, title: str = "", author: str = "") -> str:
    """Convert markdown-ish content to LaTeX source."""
    preamble = (
        "\\documentclass[12pt,a4paper]{report}\n"
        "\\usepackage[utf8]{inputenc}\n"
        "\\usepackage[T1]{fontenc}\n"
        "\\usepackage{geometry}\n"
        "\\geometry{margin=1in}\n"
        "\\usepackage{setspace}\n"
        "\\onehalfspacing\n"
        "\\usepackage{hyperref}\n"
        "\\usepackage{listings}\n"
        "\\usepackage{graphicx}\n"
        "\\usepackage{amsmath}\n\n"
    )

    if title:
        preamble += f"\\title{{{_latex_escape(title)}}}\n"
    if author:
        preamble += f"\\author{{{_latex_escape(author)}}}\n"
    preamble += "\\date{\\today}\n\n"
    preamble += "\\begin{document}\n"
    if title:
        preamble += "\\maketitle\n"
    preamble += "\\tableofcontents\n\\newpage\n\n"

    body = _convert_md_body_to_latex(text)

    return preamble + body + "\n\\end{document}\n"


def _convert_md_body_to_latex(text: str) -> str:
    """Convert markdown body text to LaTeX body."""
    lines = text.split("\n")
    output = []
    in_code = False
    in_list = False

    for line in lines:
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            if in_code:
                output.append("\\end{lstlisting}")
                in_code = False
            else:
                lang = stripped[3:].strip() or "text"
                output.append(f"\\begin{{lstlisting}}[language={lang}]")
                in_code = True
            continue

        if in_code:
            output.append(line)
            continue

        if not stripped:
            if in_list:
                output.append("\\end{itemize}")
                in_list = False
            output.append("")
            continue

        # Headings → LaTeX sections
        if stripped.startswith("#### "):
            output.append(f"\\paragraph{{{_latex_escape(stripped[5:])}}}")
        elif stripped.startswith("### "):
            output.append(f"\\subsubsection{{{_latex_escape(stripped[4:])}}}")
        elif stripped.startswith("## "):
            output.append(f"\\section{{{_latex_escape(stripped[3:])}}}")
        elif stripped.startswith("# "):
            output.append(f"\\chapter{{{_latex_escape(stripped[2:])}}}")
        # Bullets
        elif stripped.startswith("- ") or stripped.startswith("* "):
            if not in_list:
                output.append("\\begin{itemize}")
                in_list = True
            output.append(f"  \\item {_latex_escape(stripped[2:])}")
        # Regular text
        else:
            text_out = _latex_escape(stripped)
            # Convert **bold** and *italic*
            text_out = re.sub(r"\*\*([^*]+)\*\*", r"\\textbf{\1}", text_out)
            text_out = re.sub(r"\*([^*]+)\*", r"\\textit{\1}", text_out)
            text_out = re.sub(r"`([^`]+)`", r"\\texttt{\1}", text_out)
            output.append(text_out)

    if in_list:
        output.append("\\end{itemize}")
    if in_code:
        output.append("\\end{lstlisting}")

    return "\n".join(output)


def _latex_escape(text: str) -> str:
    """Escape special LaTeX characters."""
    chars = {
        "&": "\\&", "%": "\\%", "$": "\\$", "#": "\\#",
        "_": "\\_", "{": "\\{", "}": "\\}", "~": "\\textasciitilde{}",
        "^": "\\textasciicircum{}",
    }
    # Don't escape backslashes in already-escaped commands
    for char, replacement in chars.items():
        text = text.replace(char, replacement)
    return text


# ══════════════════════════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════════════════════════

def _sanitize_filename(name: str) -> str:
    """Remove unsafe characters from filenames."""
    return re.sub(r'[<>:"/\\|?*]', "_", name).strip()


# ══════════════════════════════════════════════════════════════════
#  EXPORT FULL CONVERSATION
# ══════════════════════════════════════════════════════════════════

def export_conversation(conv: dict, fmt: str = "markdown") -> str:
    """Export a full conversation (dict with 'messages') to a file.

    Supports: 'markdown', 'word', 'latex'.
    Returns the path to the saved file.
    """
    ensure_drafts_dir()
    title = conv.get("title", "Conversation")
    messages = conv.get("messages", [])

    # Build markdown content from messages
    md_parts = [f"# {title}\n"]
    for msg in messages:
        role = msg.get("role", "unknown").upper()
        content = msg.get("content", "")
        if role == "USER":
            md_parts.append(f"**You:** {content}\n")
        else:
            md_parts.append(f"**AI:** {content}\n")
        md_parts.append("---\n")

    md_content = "\n".join(md_parts)

    conv_id = conv.get("id", "export")
    if fmt == "word":
        return export_to_word(md_content, filename=f"conversation_{conv_id}.docx",
                              title=title)
    elif fmt == "latex":
        return export_to_latex(md_content, filename=f"conversation_{conv_id}.tex",
                               title=title)
    else:
        return save_draft_markdown(md_content, filename=f"conversation_{conv_id}.md")


# ══════════════════════════════════════════════════════════════════
#  BIBLIOGRAPHY / REFERENCES GENERATION
# ══════════════════════════════════════════════════════════════════

def generate_bibliography_file(papers: list[dict], style: str = "apa") -> str:
    """Generate a bibliography file from a list of paper dicts.

    Each paper dict may have: title, summary, url, added.
    Returns the file path.
    """
    ensure_drafts_dir()
    lines = [f"# References ({style.upper()} Style)\n"]

    for i, paper in enumerate(papers, 1):
        title = paper.get("title", "Untitled")
        url = paper.get("url", "")
        added = paper.get("added", "")
        year = ""
        if added:
            try:
                year = datetime.fromisoformat(added).strftime("%Y")
            except Exception:
                pass

        if style == "ieee":
            lines.append(f"[{i}] {title}. {year or '[Year]'}."
                         f"{f' Available: {url}' if url else ''}\n")
        elif style == "harvard":
            lines.append(f"- {title} ({year or 'n.d.'})."
                         f"{f' Available at: {url}' if url else ''}\n")
        else:  # apa
            lines.append(f"- {title}. ({year or 'n.d.'})."
                         f"{f' Retrieved from {url}' if url else ''}\n")

    content = "\n".join(lines)
    filepath = os.path.join(DRAFTS_DIR, f"bibliography_{style}.md")
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)
    return filepath
